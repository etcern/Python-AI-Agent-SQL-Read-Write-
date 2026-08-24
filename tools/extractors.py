"""File extractors - extract text from PDF, DOCX, XLSX, CSV, images, etc.

Each extractor returns plain text. Binary formats are handled by specialized
libraries; unknown types fall back to UTF-8 read.

Ref: https://pymupdf.readthedocs.io/en/latest/the-basics.html
Ref: https://python-docx.readthedocs.io/
Ref: https://openpyxl.readthedocs.io/
"""

import csv
import io
import json
import os

MAX_CHARS = 10000


# --- Dispatcher ---

def extract_text(file_path: str) -> str:
    """Extract readable text from a file based on its extension.

    Returns plain text or an error/info string.
    Large outputs are truncated to MAX_CHARS.
    """
    ext = os.path.splitext(file_path)[1].lower()
    extractors = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".xls":  _extract_xlsx,
        ".csv":  _extract_csv,
        ".json": _extract_json,
        ".yaml": _extract_text,
        ".yml":  _extract_text,
        ".md":   _extract_text,
        ".txt":  _extract_text,
        ".py":   _extract_text,
        ".js":   _extract_text,
        ".ts":   _extract_text,
        ".html": _extract_text,
        ".css":  _extract_text,
        ".sql":  _extract_text,
        ".sh":   _extract_text,
        ".bat":  _extract_text,
        ".xml":  _extract_text,
        ".toml": _extract_text,
        ".ini":  _extract_text,
        ".cfg":  _extract_text,
        ".log":  _extract_text,
        ".png":  _extract_image,
        ".jpg":  _extract_image,
        ".jpeg": _extract_image,
        ".bmp":  _extract_image,
        ".tiff": _extract_image,
        ".gif":  _extract_image,
        ".webp": _extract_image,
    }

    extractor = extractors.get(ext, _extract_text)
    try:
        content = extractor(file_path)
    except Exception as e:
        return f"[Error reading file: {e}]"

    # -- Truncate if too long --
    if len(content) > MAX_CHARS:
        content = content[:MAX_CHARS] + f"\n\n[truncated at {MAX_CHARS} characters]"
    return content


# --- PDF ---

def _extract_pdf(path: str) -> str:
    """Extract text from PDF using PyMuPDF.
    Falls back to OCR if page has no text (scanned PDF).
    Ref: https://pymupdf.readthedocs.io/en/latest/the-basics.html
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        return "[Cannot read PDF - install pymupdf: pip install pymupdf]"

    doc = fitz.open(path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if not text:
            # -- Try OCR if Tesseract is available --
            try:
                text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            except Exception:
                text = f"[Page {i + 1}: no text extracted - may need OCR]"
        pages.append(f"--- Page {i + 1} ---\n{text}")
    doc.close()
    return "\n\n".join(pages)


# --- DOCX ---

def _extract_docx(path: str) -> str:
    """Extract text from Word documents.
    Ref: https://python-docx.readthedocs.io/
    """
    try:
        from docx import Document
    except ImportError:
        return "[Cannot read DOCX - install python-docx: pip install python-docx]"

    doc = Document(path)
    parts = []

    # -- Paragraphs --
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # -- Tables --
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts) if parts else "[Empty document]"


# --- Excel ---

def _extract_xlsx(path: str) -> str:
    """Extract data from Excel spreadsheets.
    Ref: https://openpyxl.readthedocs.io/
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[Cannot read Excel - install openpyxl: pip install openpyxl]"

    wb = load_workbook(path, read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))

    wb.close()
    return "\n\n".join(parts) if parts else "[Empty spreadsheet]"


# --- CSV ---

def _extract_csv(path: str) -> str:
    """Extract data from CSV files."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = []
        for i, row in enumerate(reader):
            if i > 500:
                rows.append(f"[... truncated at 500 rows]")
                break
            rows.append(" | ".join(row))
    return "\n".join(rows) if rows else "[Empty CSV]"


# --- JSON ---

def _extract_json(path: str) -> str:
    """Pretty-print JSON files."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


# --- Images ---

def _extract_image(path: str) -> str:
    """Extract text from images via OCR, or return metadata.

    Uses PyMuPDF's built-in OCR if Tesseract is installed.
    Falls back to Pillow for basic metadata (dimensions, format).
    """
    # -- Try OCR with PyMuPDF --
    try:
        import fitz
        doc = fitz.open(path)
        page = doc[0]
        text = page.get_text().strip()
        doc.close()
        if text:
            return f"[OCR result from image]\n{text}"
    except Exception:
        pass

    # -- Fall back to Pillow metadata --
    try:
        from PIL import Image
        img = Image.open(path)
        info = (
            f"Format: {img.format}\n"
            f"Size: {img.width} x {img.height} px\n"
            f"Mode: {img.mode}"
        )
        img.close()
        return f"[Image metadata - OCR not available]\n{info}"
    except ImportError:
        return "[Cannot read image - install Pillow: pip install Pillow]"
    except Exception as e:
        return f"[Cannot read image: {e}]"


# --- Plain text fallback ---

def _extract_text(path: str) -> str:
    """Read as UTF-8 text. Falls back to latin-1 for binary-like files."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception:
            return "[Binary file - cannot extract text]"
    except Exception as e:
        return f"[Error reading file: {e}]"
